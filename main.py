import os
import time
import io
import logging
from datetime import datetime
from openai import OpenAI
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from dotenv import load_dotenv
from video_processor import VideoProcessor
import tempfile
import sys

# Video folder configuration
VIDEO_FOLDERS = {}

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('translator.log')
    ]
)
logger = logging.getLogger(__name__)

def check_existing_translation(file_name, service, output_folder_id):
    """Check if a translated version already exists"""
    base_without_ext = os.path.splitext(file_name)[0]
    translated_name = f"{base_without_ext}_AI_Translated.srt"
    
    logger.info(f"Checking for existing translation: {translated_name}")
    
    query = f"name='{translated_name}' and '{output_folder_id}' in parents and trashed=false"
    try:
        results = service.files().list(
            q=query,
            fields="files(id, name)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        
        if results.get('files', []):
            logger.info(f"Found existing translation: {translated_name}")
            return True
            
        logger.info("No existing translation found")
        return False
    except Exception as e:
        logger.error(f"Error checking for existing translation: {e}")
        return True  # Safer to assume it exists if we can't check

def is_video_file(mime_type):
    """Check if a file is a video based on its MIME type."""
    video_mime_types = [
        'video/mp4',
        'video/x-msvideo',  # AVI
        'video/quicktime',  # MOV
        'video/x-matroska',  # MKV
        'video/webm',
        'video/mpeg',
    ]
    return any(mime_type.startswith(vtype) for vtype in video_mime_types)

# Print startup banner
logger.info("=" * 50)
logger.info("Starting Auto Google Drive Translator")
logger.info("=" * 50)

# Load environment variables
logger.info("Loading environment variables...")
load_dotenv()

# Initialize OpenAI
logger.info("Initializing OpenAI client...")
api_key = os.getenv('OPENAI_API_KEY')
if not api_key:
    logger.error("OpenAI API key not found in environment variables!")
    sys.exit(1)

if not api_key.startswith('sk-') or len(api_key) < 20:
    logger.error("OpenAI API key appears to be invalid! Please check your .env file.")
    sys.exit(1)

try:
    client = OpenAI(
        api_key=api_key,
        base_url="https://api.openai.com/v1",
        timeout=60.0
    )
    logger.info("OpenAI client initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize OpenAI client: {e}")
    sys.exit(1)

# Setup Google Drive API
logger.info("Setting up Google Drive API...")
SCOPES = ['https://www.googleapis.com/auth/drive']
SERVICE_ACCOUNT_FILE = './credentials/credentials.json'

try:
    credentials = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    service = build('drive', 'v3', credentials=credentials)
    logger.info("Google Drive API setup completed successfully")
except Exception as e:
    logger.error(f"Failed to initialize Google Drive API: {e}")
    sys.exit(1)

def list_files(folder_id, suffix=None):
    """List all files in the specified Google Drive folder."""
    try:
        query = f"'{folder_id}' in parents and trashed=false"
        if suffix:
            query += f" and name ends with '{suffix}'"
        
        results = service.files().list(
            q=query,
            pageSize=100,
            fields="files(id, name, mimeType)",
            orderBy="name"
        ).execute()
        
        folder_name = "Unknown Folder"
        folder_names = {
            os.getenv('VIDEO_DUTCH_FOLDER_ID'): "Dutch Videos",
            os.getenv('VIDEO_GERMAN_FOLDER_ID'): "German Videos",
            os.getenv('VIDEO_ENGLISH_FOLDER_ID'): "English Videos",
            os.getenv('VIDEO_OTHER_FOLDER_ID'): "Other Videos",
            os.getenv('OUTPUT_FOLDER_ID'): "Output",
            os.getenv('TEXT_TRANSLATION_FOLDER_ID'): "Translation Queue"
        }
        folder_name = folder_names.get(folder_id, folder_id)
        files = results.get('files', [])
        logger.info(f"Found {len(files)} files in {folder_name}")
        return files
    except Exception as e:
        logger.error(f"Error listing files in {folder_name}: {e}")
        return []

def download_file(file_id, file_name):
    """Download a file from Google Drive."""
    logger.info(f"Downloading file: {file_name}")
    data_folder = os.path.join(os.getcwd(), "Data")
    os.makedirs(data_folder, exist_ok=True)
    file_path = os.path.join(data_folder, file_name)
    
    try:
        request = service.files().get_media(fileId=file_id)
        with open(file_path, 'wb') as fh:
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
                logger.info(f"Download progress: {int(status.progress() * 100)}%")
        logger.info(f"Successfully downloaded {file_name}")
        return file_path
    except Exception as e:
        logger.error(f"Error downloading file {file_name}: {e}")
        if os.path.exists(file_path):
            os.remove(file_path)
        return None

def translate_text(text):
    """Translate text using OpenAI's API."""
    try:
        logger.info("Starting text translation")
        lines = text.split("\n")
        chunks = [lines[i:i+125] for i in range(0, len(lines), 125)]
        translated_chunks = []

        for i, chunk in enumerate(chunks, 1):
            logger.info(f"Translating chunk {i}/{len(chunks)}")
            chunk_text = "\n".join(chunk)
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": os.getenv('SYSTEM_PROMPT', "You are a helpful translation assistant.")},
                    {"role": "user", "content": chunk_text}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            translated_text = response.choices[0].message.content
            translated_text = translated_text.replace("```plaintext", "").replace("```", "")
            translated_chunks.append(translated_text)
            time.sleep(1)  # Rate limiting

        logger.info("Translation completed successfully")
        return "\n".join(translated_chunks)
    except Exception as e:
        logger.error(f"Translation error: {e}")
        return None

def translate_file(file_path, mime_type):
    """Translate file content based on mime type."""
    logger.info(f"Starting file translation for {file_path}")
    try:
        file_dir, file_name = os.path.split(file_path)
        file_name_without_ext, file_ext = os.path.splitext(file_name)
        new_file_name = f"{file_name_without_ext}_AI_Translated{file_ext}"
        new_file_path = os.path.join(file_dir, new_file_name)

        if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            logger.info("Processing Word document")
            doc = Document(file_path)
            new_doc = Document()
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    translated_text = translate_text(paragraph.text)
                    if translated_text:
                        new_doc.add_paragraph(translated_text)
            new_doc.save(new_file_path)

        elif mime_type == 'application/pdf':
            logger.info("Processing PDF document")
            reader = PdfReader(file_path)
            writer = PdfWriter()
            
            for page_num, page in enumerate(reader.pages, 1):
                logger.info(f"Processing PDF page {page_num}/{len(reader.pages)}")
                text = page.extract_text()
                if text.strip():
                    translated_text = translate_text(text)
                    if translated_text:
                        packet = io.BytesIO()
                        c = canvas.Canvas(packet)
                        c.drawString(100, 700, translated_text)
                        c.save()
                        packet.seek(0)
                        new_page = PdfReader(packet).pages[0]
                        writer.add_page(new_page)

            with open(new_file_path, 'wb') as output_file:
                writer.write(output_file)

        else:  # Text-based files
            logger.info("Processing text-based file")
            translated_lines = []
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                translated_text = translate_text(text)
                if translated_text is None:
                    logger.error("Translation failed")
                    return None
                    
                with open(new_file_path, 'w', encoding='utf-8') as output_file:
                    output_file.write(translated_text)
                    
        if os.path.exists(new_file_path):
            logger.info(f"Successfully translated file to {new_file_path}")
            return new_file_path
        else:
            logger.error(f"Translation file not created: {new_file_path}")
            return None
            
    except Exception as e:
        logger.error(f"Error translating file: {e}")
        if 'new_file_path' in locals() and os.path.exists(new_file_path):
            try:
                os.remove(new_file_path)
                logger.info(f"Cleaned up partial translation file: {new_file_path}")
            except Exception as cleanup_error:
                logger.error(f"Error cleaning up translation file: {cleanup_error}")
        return None

def upload_file(file_path, folder_id, relative_path=None):
    """Upload file to Google Drive."""
    logger.info(f"Uploading file: {file_path}")
    try:
        current_folder_id = folder_id
        if relative_path:
            for folder in relative_path:
                current_folder_id = create_or_get_folder(folder, current_folder_id)

        file_name = os.path.basename(file_path)
        file_metadata = {
            'name': file_name,
            'parents': [current_folder_id]
        }
        
        media = MediaFileUpload(file_path, resumable=True)
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()
        
        logger.info(f"Successfully uploaded {file_name}")
        return file.get('id')
    except Exception as e:
        logger.error(f"Error uploading file: {e}")
        return None

def process_file(file_id, file_name, mime_type, folder_path=None):
    """Process a single file."""
    try:
        # Check if translation already exists
        output_folder_id = os.getenv('OUTPUT_FOLDER_ID')
        if check_existing_translation(file_name, service, output_folder_id):
            logger.info(f"Skipping {file_name} - translation already exists")
            return True

        logger.info(f"Processing file: {file_name}")
        
        # Download and process the file
        file_path = download_file(file_id, file_name)
        if not file_path:
            return False

        translated_path = translate_file(file_path, mime_type)
        if not translated_path:
            return False

        if upload_file(translated_path, os.getenv('OUTPUT_FOLDER_ID'), folder_path):
            logger.info(f"Successfully processed {file_name}")
            return True
        return False
    except Exception as e:
        logger.error(f"Error processing file {file_name}: {e}")
        logger.error("Error details:", exc_info=True)
        return False
    finally:
        # Cleanup
        try:
            if 'file_path' in locals() and file_path and os.path.exists(file_path):
                os.remove(file_path)
            if 'translated_path' in locals() and translated_path and os.path.exists(translated_path):
                os.remove(translated_path)
        except Exception as e:
            logger.error(f"Error during cleanup: {e}")

def load_environment():
    """Load and validate environment variables."""
    if not load_dotenv():
        logger.error("Failed to load .env file")
        return False
        
    global VIDEO_FOLDERS
    
    # Initialize video folders after environment is loaded
    VIDEO_FOLDERS = {
        "Dutch": {
            "code": "nl",
            "folder_id": os.environ.get("VIDEO_DUTCH_FOLDER_ID")
        },
        "German": {
            "code": "de", 
            "folder_id": os.environ.get("VIDEO_GERMAN_FOLDER_ID")
        },
        "English": {
            "code": "en",
            "folder_id": os.environ.get("VIDEO_ENGLISH_FOLDER_ID")
        },
        "Other": {
            "code": None,
            "folder_id": os.environ.get("VIDEO_OTHER_FOLDER_ID")
        }
    }

    # Validate all required environment variables
    required_vars = {
        'VIDEO_DUTCH_FOLDER_ID': ('Dutch videos', VIDEO_FOLDERS["Dutch"]["folder_id"]),
        'VIDEO_GERMAN_FOLDER_ID': ('German videos', VIDEO_FOLDERS["German"]["folder_id"]),
        'VIDEO_ENGLISH_FOLDER_ID': ('English videos', VIDEO_FOLDERS["English"]["folder_id"]),
        'VIDEO_OTHER_FOLDER_ID': ('Other videos', VIDEO_FOLDERS["Other"]["folder_id"]),
        'TEXT_TRANSLATION_FOLDER_ID': ('Text translations', os.getenv('TEXT_TRANSLATION_FOLDER_ID')),
        'OUTPUT_FOLDER_ID': ('Output', os.getenv('OUTPUT_FOLDER_ID')),
        'OPENAI_API_KEY': ('OpenAI API', os.getenv('OPENAI_API_KEY'))
    }

    missing_vars = []
    
    for var, (description, value) in required_vars.items():
        if not value:
            missing_vars.append(description)
            logger.warning(f"No folder ID configured for {description}")
        else:
            logger.info(f"Loaded {description} folder ID: {value[:10]}...")

    if missing_vars:
        logger.error(f"Missing required environment variables for: {', '.join(missing_vars)}")
        return False

    return True

def main():
    logger.info("==================================================")
    logger.info("Starting Auto Google Drive Translator")
    logger.info("==================================================")
    
    logger.info("Loading environment variables...")
    if not load_environment():
        logger.error("Failed to load required environment variables")
        return

    logger.info("Starting main processing loop")
    processed_files = set()  # Keep track of processed files

    while True:
        try:
            # Process text folder
            text_folder_id = os.getenv('TEXT_TRANSLATION_FOLDER_ID')
            if text_folder_id:
                logger.info("Checking text translation folder")
                files = list_files(text_folder_id)
                
                for file in files:
                    file_id = file['id']
                    file_name = file['name']
                    
                    # Skip if already processed in this session or has _AI_Translated suffix
                    if (file_id in processed_files or 
                        "_AI_Translated" in file_name or 
                        file_name.lower().endswith('_ai_translated.srt')):
                        logger.info(f"Skipping already processed file: {file_name}")
                        continue
                        
                    if process_file(file_id, file_name, file['mimeType']):
                        processed_files.add(file_id)

            # Process videos in language-specific folders
            logger.info("Checking video folders...")
            video_processor = None  # Initialize only if needed
            
            for language, folder_info in VIDEO_FOLDERS.items():
                folder_id = folder_info.get("folder_id")
                if not folder_id:
                    logger.warning(f"No folder ID configured for {language} videos")
                    continue
                    
                logger.info(f"Checking {language} video folder")
                files = list_files(folder_id)
                
                for file in files:
                    file_id = file['id']
                    file_name = file['name']
                    
                    if (file_id not in processed_files and 
                        is_video_file(file['mimeType'])):
                        logger.info(f"Found new video to process: {file_name} (type: {file['mimeType']})")
                        
                        if video_processor is None:
                            video_processor = VideoProcessor()
                            
                        # Pass folder_info instead of language string
                        if video_processor.process_video(file_id, file_name, folder_info, service):
                            logger.info(f"Successfully processed video: {file_name}")
                            processed_files.add(file_id)
                        else:
                            logger.error(f"Failed to process video: {file_name}")

            logger.info("Waiting for new files...")
            time.sleep(60)

        except Exception as e:
            logger.error(f"Error in main loop: {e}")
            logger.error("Error details:", exc_info=True)
            time.sleep(60)

if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        logger.info("Shutting down gracefully...")
        sys.exit(0)
    except Exception as e:
        logger.error(f"Fatal error: {e}")
        sys.exit(1)
